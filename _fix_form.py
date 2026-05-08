"""修复案例信息表"""
from docx import Document

path = '/Users/chiu/Desktop/创ai/动物习作AI智能助教/案例信息表.docx'
doc = Document(path)

full_text = (
    '本案例聚焦四年级语文动物类习作教学中\u201c个性化辅导不足\u201d的核心痛点，'
    '借助国产DeepSeek大模型自主开发《动物习作AI智能助教》教育智能体。'
    '该工具包含写作闯关、好词好句墙、写作分析建议、互动问答、范文展示五大功能模块，'
    '采用Flask本地Web架构，无需部署服务器。'
    '经惠州市富民小学四年级4个班级试用，学生习作字数平均提升32%，'
    '写作兴趣显著提高，教师批改效率提升约60%。'
    '本工具采用MIT开源协议发布。'
)

for table in doc.tables:
    for row in table.rows:
        text = ' '.join([c.text for c in row.cells])
        if '案例内容简介' in text:
            cell = row.cells[1]
            content = cell.text.strip()
            print(f'当前长度: {len(content)} 字')
            for p in cell.paragraphs:
                for r in p.runs:
                    r.text = ''
                if p.runs:
                    p.runs[0].text = full_text
            print(f'更新后: {len(full_text)} 字')

doc.save(path)
print('案例信息表已更新')
