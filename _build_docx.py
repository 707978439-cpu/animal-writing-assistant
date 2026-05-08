import re
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()
section = doc.sections[0]
section.page_width = Cm(21.0)
section.page_height = Cm(29.7)
section.top_margin = Cm(2.54)
section.bottom_margin = Cm(2.54)
section.left_margin = Cm(3.17)
section.right_margin = Cm(3.17)

def set_font(run, fn, sz, bold=False):
    run.font.size = Pt(sz)
    run.font.bold = bold
    run.font.name = fn
    rPr = run._r.get_or_add_rPr()
    rf = rPr.find(qn('w:rFonts'))
    if rf is None:
        rf = OxmlElement('w:rFonts')
        rPr.append(rf)
    rf.set(qn('w:eastAsia'), fn)
    rf.set(qn('w:ascii'), fn)
    rf.set(qn('w:hAnsi'), fn)

def add_para(text, fn='仿宋_GB2312', sz=16, bold=False, align=None, indent=True):
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    if indent:
        pPr = p._p.get_or_add_pPr()
        ind = pPr.find(qn('w:ind'))
        if ind is None:
            ind = OxmlElement('w:ind')
            pPr.append(ind)
        ind.set(qn('w:firstLine'), str(int(16 * 20 * 2)))
    sp = p._p.get_or_add_pPr()
    spc = sp.find(qn('w:spacing'))
    if spc is None:
        spc = OxmlElement('w:spacing')
        sp.append(spc)
    spc.set(qn('w:line'), '560')
    spc.set(qn('w:lineRule'), 'exact')
    run = p.add_run(text)
    set_font(run, fn, sz, bold)
    return p

base_path = '/Users/chiu/Desktop/创ai/动物习作AI智能助教'

add_para('开发与应用报告', '方正小标宋简体', 18, align=WD_ALIGN_PARAGRAPH.CENTER, indent=False)
add_para('')

with open(base_path + '/development_report_content.md', 'r') as f:
    content = f.read()

for line in content.split('\n'):
    line = line.strip()
    if not line:
        continue
    if line.startswith('### '):
        add_para(line[4:], '楷体_GB2312', 16)
    elif line.startswith('## '):
        add_para(line[3:], '黑体', 16)
    elif line.startswith('# '):
        pass
    elif line.startswith('```'):
        pass
    elif line.startswith('**') and '：' in line:
        add_para(line.strip('*'), '楷体_GB2312', 16)
    else:
        add_para(line)

output = base_path + '/开发与应用报告.docx'
doc.save(output)
print('DOCX saved to:', output)
full_text = ''.join([p.text for p in doc.paragraphs])
chinese = len(re.findall(r'[\u4e00-\u9fff]', full_text))
print('Chinese chars:', chinese)
